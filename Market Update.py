from urllib.request import Request, urlopen
from bs4 import BeautifulSoup as soup
from tabulate import tabulate
import docx
from datetime import datetime

doc = docx.Document()
now = datetime.now() 

date_time = now.strftime("%m.%d.%Y")
url = "https://finviz.com/groups.ashx?g=sector&v=110&o=-change"
req = Request (url, headers ={'User-Agent': 'Mozzila/5.0'})
webpage = urlopen(req).read()
source = soup(webpage, 'html.parser')
p_tags = source.find_all('td', {'class' :['body-table']}) # <td align="right" class="body-table" height="10">11</td>

trill = []

for each in p_tags:
    x = (str(each.get_text()))
    trill.append(x)

mydata = [(trill[0],  trill[1], trill[9], trill[5], trill[6], trill[4], trill[8], trill[3]), #rank, sector, $change, current p/e, forward p/e, div yield, short interest, market cap
          (trill[11], trill[12], trill[20], trill[16], trill[17], trill[15], trill[19], trill[14]),
          (trill[22], trill[23], trill[31], trill[27], trill[28], trill[26], trill[30], trill[25]),
          (trill[33], trill[34], trill[42], trill[38], trill[39], trill[37], trill[41], trill[36]),
          (trill[44], trill[45], trill[53], trill[49], trill[50], trill[48], trill[52], trill[47]),
          (trill[55], trill[56], trill[64], trill[60], trill[61], trill[59], trill[63], trill[58]),
          (trill[66], trill[67], trill[75], trill[71], trill[72], trill[70], trill[74], trill[69]),
          (trill[77], trill[78], trill[86], trill[82], trill[83], trill[81], trill[85], trill[80]),
          (trill[88], trill[89], trill[97], trill[93], trill[94], trill[92], trill[96], trill[91]),
          (trill[99], trill[100], trill[108], trill[104], trill[105], trill[103], trill[107], trill[102]),
          (trill[110], trill[111], trill[119], trill[115], trill[116], trill[114], trill[118], trill[113])
          ]

headers = ['Rank', 'Sector', '% Change','Current P/E', 'Forward P/E', 'Div Yield', 'Short Interest', 'Mrk Cap']

Happy = (tabulate(mydata, headers = headers )) #,tablefmt ='grid' add for grid lines


paraObject = doc.add_paragraph(print(Happy))
paraObject = doc.add_paragraph(Happy)
doc.save('Sector'+date_time)

print('doc saved')
